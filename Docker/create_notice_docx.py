from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUT_PATH = BASE_DIR / "Notice_Docker_NVIDIA_PaginaVox.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(35, 35, 35)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(10)

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_footer(doc):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PaginaVox Docker NVIDIA")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_summary_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_width(table)
    set_cell_margins(table)
    headers = table.rows[0].cells
    headers[0].text = "Element"
    headers[1].text = "Role"
    for cell in headers:
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    rows = [
        ("Docker/Dockerfile", "Image CUDA/Python avec Torch, Qwen TTS, Whisper et Gradio."),
        ("Docker/docker-compose.yml", "Lancement GPU, ports, volumes et cache Hugging Face."),
        ("Docker/build_windows.bat", "Construction de l'image sous Windows."),
        ("Docker/run_windows.bat", "Lancement de l'application sous Windows."),
        ("Docker/test_gpu.bat / .sh", "Test rapide de l'acces NVIDIA dans Docker."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right


def build_doc():
    doc = Document()
    setup_styles(doc)
    add_footer(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Notice Docker NVIDIA - PaginaVox")
    subtitle = doc.add_paragraph()
    subtitle.add_run("Guide d'installation, de build et de lancement avec GPU NVIDIA").italic = True

    doc.add_heading("Objectif", level=1)
    doc.add_paragraph(
        "Ce document explique comment utiliser le dossier Docker/ pour construire et lancer "
        "PaginaVox dans Docker avec acceleration GPU NVIDIA. L'interface Gradio est lancee par defaut."
    )

    doc.add_heading("Fichiers fournis", level=1)
    add_summary_table(doc)

    doc.add_heading("Prerequis", level=1)
    add_bullets(
        doc,
        [
            "Un PC avec carte NVIDIA compatible CUDA.",
            "Des pilotes NVIDIA recents installes sur la machine hote.",
            "Docker Desktop avec backend WSL2 sur Windows, ou Docker Engine sur Linux.",
            "Une connexion internet pour construire l'image et telecharger les modeles au premier lancement.",
            "De l'espace disque disponible : l'image, Torch et les modeles peuvent prendre plusieurs Go.",
        ],
    )

    doc.add_heading("Verification GPU", level=1)
    doc.add_paragraph("Avant de construire PaginaVox, teste que Docker voit bien la carte NVIDIA.")
    doc.add_paragraph("Windows :")
    add_code(doc, r".\Docker\test_gpu.bat")
    doc.add_paragraph("Linux :")
    add_code(doc, "bash Docker/test_gpu.sh")
    doc.add_paragraph("Le test doit afficher le tableau nvidia-smi. Si ce n'est pas le cas, corrige Docker/NVIDIA avant de continuer.")

    doc.add_heading("Construire l'image", level=1)
    doc.add_paragraph("Windows :")
    add_code(doc, r".\Docker\build_windows.bat")
    doc.add_paragraph("Linux :")
    add_code(doc, "bash Docker/build_linux_macos.sh")
    doc.add_paragraph("L'image creee s'appelle paginavox:gpu.")

    doc.add_heading("Lancer l'interface Gradio", level=1)
    doc.add_paragraph("Windows :")
    add_code(doc, r".\Docker\run_windows.bat")
    doc.add_paragraph("Linux :")
    add_code(doc, "bash Docker/run_linux_macos.sh")
    doc.add_paragraph("Ouvre ensuite l'adresse suivante dans le navigateur :")
    add_code(doc, "http://localhost:7860")

    doc.add_heading("Lancer le mode terminal", level=1)
    doc.add_paragraph("Si tu veux retrouver le mode question/reponse en terminal dans Docker :")
    add_code(doc, "docker compose -f Docker/docker-compose.yml run --rm paginavox python3 main.py")

    doc.add_heading("Dossiers et donnees", level=1)
    add_bullets(
        doc,
        [
            "audio/ : audios de reference pour le clonage.",
            "txt/ : textes et transcriptions Whisper.",
            "output/ : fichiers WAV generes.",
            "profiles/ : profils de voix clones.",
            "paginavox-hf-cache : volume Docker qui garde les modeles Hugging Face telecharges.",
        ],
    )

    doc.add_heading("Notes Windows", level=1)
    add_steps(
        doc,
        [
            "Installe Docker Desktop.",
            "Active le backend WSL2 dans Docker Desktop.",
            "Installe ou mets a jour le pilote NVIDIA Windows.",
            "Lance le test GPU du dossier Docker.",
            "Construis puis lance PaginaVox.",
        ],
    )

    doc.add_heading("Note macOS", level=1)
    doc.add_paragraph(
        "Docker Desktop sur macOS ne donne pas acces a une carte NVIDIA CUDA. "
        "Cette configuration est donc prevue pour Windows avec WSL2/Docker Desktop "
        "ou Linux avec NVIDIA Container Toolkit."
    )

    doc.add_heading("Depannage", level=1)
    add_bullets(
        doc,
        [
            "Erreur GPU introuvable : verifier les pilotes NVIDIA, WSL2 et le support GPU Docker.",
            "Build tres long : normal, les dependances IA sont lourdes.",
            "Premier lancement long : les modeles Qwen ou Whisper peuvent se telecharger.",
            "Memoire GPU insuffisante : fermer les autres applications GPU ou utiliser une machine plus puissante.",
            "Port 7860 deja utilise : changer le mapping dans Docker/docker-compose.yml.",
        ],
    )

    doc.add_heading("Sources officielles", level=1)
    add_bullets(
        doc,
        [
            "Docker GPU Desktop Windows : https://docs.docker.com/desktop/features/gpu/",
            "Docker Compose GPU support : https://docs.docker.com/compose/gpu-support/",
            "NVIDIA Container Toolkit : https://docs.nvidia.com/datacenter/cloud-native/container-toolkit/latest/",
            "CUDA on WSL : https://docs.nvidia.com/cuda/wsl-user-guide/index.html",
        ],
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_doc()
